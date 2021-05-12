"""
    Author: Ibrahim, Ibrahim Opeyemi
        Email: IbrahimIbrahimOpeyemi@gmail.com
        Phone: 08107321115
    Version: 1.2
    This script copies articles from Ajol. It works with ajol and
         probably any aggregator sites using exactly the same HTML template
    ---------------usage-------------
    1. Ensure the directory the script resides is empty or at least
        doesn't have and word documents
    2. launch the script and run in terminal 
        copy and paste the journal archive url in the terminal prompt
    3. Input the position of the issue to start copying from. If
    4. When the "All issues listed in the page have been saved successfully"
        message appears, move the generated word documents to the where you'd like
        to store them
    NB: The scripts overwrite existing file when name conflict occur
    tabs and line feed are usually present in the generated files.
    this will ease finally formatting of these documents

    ---------------dependencies-----------
    1.  BeautifulSoup4==4.9.1
    2.  requests==2.24.0


"""
from bs4 import BeautifulSoup
import requests
import unicodedata
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

issues_listing_url = input("Input Issues listing URL ('Archive'): ").strip()

try:
    start_at = input(
        "Input the start Issue position on the page(default=1): ").strip()

    start_at = int(start_at)
    print(f"Starting at issue no {start_at}")
    start_at = start_at - 1
except ValueError:
    print("Invalid position. Only integer allowed")
    start_at = -1


def fetch_biography(url):
    r_text = get_reponse_text(url)
    if r_text != False:
        soup = create_bsoup(r_text)
    try:
        author_bio_divs = soup.find_all(
            'div', class_="item author_bios")
        if len(author_bio_divs) > 0:
            author_bio = "".join([a.getText() for a in author_bio_divs])
            author_bio = f"({author_bio})"
            author_bio = author_bio.replace('\t', " ")
            return author_bio
        else:
            author_bio = ""
            return author_bio
    except IndexError:
        author_bio = ""
        return author_bio


def get_journal_name(journal_name_bsoup):
    try:
        journal_name = journal_name_bsoup.select(
            '#headerNavigationContainer > div > div > div > div:nth-child(1) > a')
        journal_name = journal_name[0].getText().strip()
    except IndexError:
        journal_name = ""
    return journal_name


def get_issue_url(issue_listing_url):
    try:
        issue_response_text = get_reponse_text(issue_listing_url)

        issues_soup = create_bsoup(issue_response_text)

        issues_url_list_temp = issues_soup.select(
            '#pkp_content_main > div.page.page_issue_archive > ul > li > div > a')
        # print(type(issues_url_list_temp))
        # print(len(issues_url_list_temp))
        issues_url_list = [a.get('href').strip()
                           for a in issues_url_list_temp]

        return issues_url_list

    except IndexError:
        issues_url_list = False
        return issues_url_list


def get_volume(bsoup):
    try:
        get_volume = bsoup.select(
            "#pkp_content_main > div.page.page_issue > nav > ol > li.current")
        volume = get_volume[0].getText().strip()
        return volume
    except IndexError:
        volume = ". No volume Number Not Found! "
        return volume


def get_page_number(a):
    try:
        page_number = a.select(
            '.obj_article_summary .pages')[0].getText().strip().replace("\t", "")
        return page_number

    except IndexError:
        page_number = ". Page Number Not Found! "
        return page_number


def get_title(a):
    try:
        title = a.select('.obj_article_summary .title>a')[
            0].getText().strip().replace("\t", "")
        return title
    except IndexError:
        title = ". Title Not Found! "
        return title


def get_article_url(a):
    try:
        article_url = a.select('.obj_article_summary .title>a')[0].get('href')
        return article_url
    except IndexError:
        article_url = ". Article link Not Found! "
        return article_url


def get_authors(a):
    try:
        authors = a.select('.meta .authors')[
            0].getText().strip().replace("\t", "")
        return authors
    except IndexError:
        authors = ". Authors Not Found! "
        return authors


def fetch_abstract(url):
    r_text = get_reponse_text(url)
    if r_text != False:
        abstract_bsoup = create_bsoup(r_text)
        try:
            abstract = abstract_bsoup.select(
                "#pkp_content_main > div.page.page_article > article > div > div.main_entry > div.item.abstract")
            abstract_temp = abstract[0].getText()
            abstract = abstract_temp
            abstract = abstract.replace("\t", " ")
            # print("First Conditional at work")
            return abstract
        except IndexError:
            try:
                abstract = abstract_bsoup.select(
                    "#pkp_content_main > div.page.page_article > article > div > div.main_entry > div.item.abstract p")
                abstract_temp = "".join([a.getText() for a in abstract])
                abstract = abstract_temp.replace("\t", " ")
                # print("Second Conditional at work!")
                return abstract
            except IndexError:
                abstract = ". No Abstract Found! "

        return abstract


def control_char_remover(text):
    treated = "".join([char for char in text if
                       unicodedata.category(char)[0] != "C"])
    return treated


def get_reponse_text(url):
    response = requests.get(url)
    if response.status_code == 200:
        response_text = response.text
        response_text = response_text.encode('utf-8', errors='replace')
        response_text = response_text.decode('utf-8')
        response_text = control_char_remover(response_text)
        return response_text
    else:
        return False


def create_bsoup(response_text):
    bsoup = BeautifulSoup(response_text, 'html.parser')
    return bsoup


issues_list = get_issue_url(issues_listing_url)
issues_list = issues_list[start_at:]
issue_i = 0
for issue in issues_list:
    response_text = get_reponse_text(issue)
    # implementation
    if (response_text != False) and (start_at != -1):
        bsoup = create_bsoup(response_text)
        articles = bsoup.select(
            ".obj_article_summary"
        )
        usable_volume = get_volume(bsoup)
        # create blank document

        mydoc = docx.Document()
        # Save working document name. Prefix journal name +vol+issue details
        journal_name = get_journal_name(bsoup)
        file_name = f"{journal_name}_{usable_volume}.docx"
        file_name = file_name.replace(":", "")
        print(f"creating file: {file_name}")
        # create style
        style = mydoc.styles['Normal']
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        mydoc.save(file_name)

        i = 0
        for a in articles:
            # use this list to skip articles with character not
            # compatible with XML(unicode or ASCII)
            # empty when done with the volume
            # if i in []:
            #     i += 1
            #     continue
            usable_title = get_title(a)
            usable_authors = get_authors(a)
            usable_page_number = get_page_number(a)
            usable_article_url = get_article_url(a)
            usable_abstract = fetch_abstract(usable_article_url)
            usable_biography = fetch_biography(usable_article_url)

            mydoc = docx.Document(file_name)
            paragraph = mydoc.add_paragraph(
                f"{usable_authors}. {usable_biography}. ")
            paragraph.add_run(f"{usable_title}. ").bold = True
            paragraph.add_run(
                f"{usable_volume}: {usable_page_number}.  {usable_abstract} ")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            mydoc.add_paragraph("")
            mydoc.save(file_name)
            i += 1
            print(f"Articles Processed: {i}")
        issue_i += 1
        print(f"No of issues saved: {issue_i}")

print("All issues listed in the page have been saved successfully 😊")
